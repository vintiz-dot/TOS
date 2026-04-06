from docx import Document
doc = Document('Copy of Science Moreland Lesson Plan Template .docx')
all_text = []
for p in doc.paragraphs:
    all_text.append(p.text)
all_text.append("\n=== TABLES ===")
for i, t in enumerate(doc.tables):
    all_text.append(f"\n--- Table {i+1} ---")
    for row in t.rows:
        cells = [c.text for c in row.cells]
        all_text.append(" | ".join(cells))
result = "\n".join(all_text)
with open('docx_text.txt', 'w', encoding='utf-8') as f:
    f.write(result)
print(f"DOCX extracted: {len(result)} chars")
